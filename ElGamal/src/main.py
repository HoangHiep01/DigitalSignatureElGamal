from ElGamal import ElGamal



def readFile(fileName):

	if ".txt" in fileName:
		file = open(fileName, 'r')
		message = file.read()
		file.close()
		return message
	elif ".docx" in fileName or ".doc" in fileName:
		import docx
		doc = docx.Document(fileName)
		message = ""
		message = [message + str(doc.paragraphs[i].text) for i in range(len(doc.paragraphs))]
		message = "".join(message)
		doc.save(fileName)
		return message
	elif ".pdf" in fileName:
		import PyPDF2
		pdfFileObj = open(fileName, 'rb')
		pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
		# print(pdfReader.numPages)
		# pageObj = pdfReader.getPage(0)
		# message = pageObj.extractText()
		message = ""
		message = [message + str(pdfReader.getPage(i).extractText()) for i in range(pdfReader.numPages)]
		message = "".join(message)
		pdfFileObj.close()
		return message
	else:
		print ("Invalid File")

def main():

	fileName = ("../test/test.pdf")
	message = readFile(fileName)

	print (message)
	
	obj = ElGamal()
	obj.setHashed(message)
	obj.autoFillP()
	obj.autoFillA()
	obj.autoFillAlpha()
	obj.autoFillBeta()
	obj.display()
	obj.findRandomPrivateKey()
	obj.SignatureGeneration()
	print (obj.SignatureVerification(readFile(fileName)))

if __name__ == "__main__":
	main()