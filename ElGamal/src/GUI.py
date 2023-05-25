from tkinter import Tk, Entry, Label, Button, Text, END, LabelFrame, messagebox, filedialog
from ElGamal import ElGamal
import docx
import PyPDF2

root = Tk()
root.iconbitmap("../icon/logo.ico")
root.title("Chữ ký điện tử ElGamal")
root.geometry("1250x500")

obj = ElGamal()
obj2 = ElGamal()

def generateKey():
	obj.autoFillP()
	obj.autoFillA()
	obj.autoFillAlpha()
	obj.autoFillBeta()
	obj.findRandomPrivateKey()

	entryParameter.delete(0, END)
	entryPubKey.delete(0, END)
	entryPriKey.delete(0, END)
	entryPriK.delete(0, END)
	entryParameter.insert(0, str(obj.getParameter()))
	entryPubKey.insert(0, str(obj.getPubKey()))
	entryPriKey.insert(0, str(obj.getPriKey()))
	entryPriK.insert(0, str(obj.getPriK()))


def messageBox(title, sign):

	if sign == True:
		message = "Successful"
		messagebox.showinfo(title, message)
	elif sign == False:
		message = "Failed"
		messagebox.showerror(title, message)

def readFile(filedir):

	if ".txt" in filedir:
		file = open(filedir, 'r', encoding = 'utf-8')
		message = file.read()
		file.close()
		return message
	elif ".docx" in filedir or ".doc" in filedir:
		doc = docx.Document(filedir)
		message = ""
		message = [message + str(doc.paragraphs[i].text) for i in range(len(doc.paragraphs))]
		message = "".join(message)
		doc.save(filedir)
		return message
	elif ".pdf" in filedir:
		pdfFileObj = open(filedir, 'rb')
		pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
		# print(pdfReader.numPages)
		# pageObj = pdfReader.getPage(0)
		# message = pageObj.extractText()
		message = ""
		message = [message + str(pdfReader.getPage(i).extractText()) for i in range(pdfReader.numPages)]
		message = "".join(message)
		pdfFileObj.close()
		return message
	else:
		 messageBox("Erorr", False)

def openFile(sign):
	root.filedir = filedialog.askopenfilename(initialdir="../test/", title="Choose File", filetypes=(("All Files","*.*"),("doc","*.doc"),("docx","*.docx"),("text","*.txt"),("pdf","*.pdf")))
	message = "Erorr!"
	message = readFile(root.filedir)
	if sign == "MS":
		entryMess.delete(0.0,END)
		entryMess.insert(0.0, message)
		obj.setHashed(message)
	elif sign == "MR":
		entryMess2.delete(0.0,END)
		entryMess2.insert(0.0, message)
		obj2.setHashed(message)
	elif sign == "DS":
		message = message.split()
		message = [int(char) for char in message]
		ans = ""
		for num in message:
			ans += chr(num)
		entrySignature2.delete(0.0, END)
		entrySignature2.insert(0.0, ans)
		

def signatureGeneration():
	message = entryMess.get(0.0,END)
	obj.setHashed(message)
	obj.SignatureGeneration()
	entrySignature.delete(0.0,END)
	entrySignature.insert(0.0, obj.getDigitalSignature())
	saveDigSig = open("../digitalsignature/signature.txt", "w")
	saveDigSig.write(str(obj.sign1) + "\n")
	for num in obj.sign2:
		saveDigSig.write(str(num) + " ")
	saveDigSig.close()
	messageBox("Signature Generation", True)

def enterPublicKey():
	obj2.setP(int(entryPubKeyP.get()))
	obj2.setAlpha(int(entryPubKeyAlpha.get()))
	obj2.setBeta(int(entryPubKeyBeta.get()))
	messageBox("Enter Public Key", True)

def signatureVerification():

	readDigSig = open("../digitalsignature/signature.txt", "r")
	allchar = readDigSig.readlines()
	num = int(allchar[0])
	listNum = allchar[1].split()
	listNum = [int(char) for char in listNum]
	readDigSig.close()
	obj2.setSign1(num)
	obj2.setSign2(listNum)
	ans = obj2.SignatureVerification(str(entryMess2.get(0.0, END)))
	messageBox("Signature Verification", ans)
# Tạo giao diện chính

frame1 = LabelFrame(root, text="Key")
labelParameter = Label(frame1, text="Parameter (p, alpha)")
entryParameter = Entry(frame1, width=30,borderwidth=3)
labelPubKey = Label(frame1, text="Puclic key (beta)")
entryPubKey = Entry(frame1, width=30,borderwidth=3)
labelPriKey = Label(frame1, text="Private key (a)")
entryPriKey = Entry(frame1, width=30,borderwidth=3)
labelPriK = Label(frame1, text="Private k (k)")
entryPriK = Entry(frame1, width=30,borderwidth=3)
generateKeyBtn = Button(frame1, text="Genegate Key", command= generateKey)

frame2 = LabelFrame(root, text="Message Sent")
labelMess = Label(frame2, text="Message Sent")
entryMess = Text(frame2,height=10,width=30,borderwidth=3)
# entryMess = Entry(frame2, width=30,borderwidth=3)
labelSignature = Label(frame2, text="Digital Signature Sent")
entrySignature = Text(frame2,height=10,width=30, borderwidth=3)
# entrySignature = Entry(frame2, width=30,borderwidth=3)
generateSignatureBtn = Button(frame2, text="Signature Generation", command= signatureGeneration)
# enterBtn = Button(frame2, text="Enter")
chooseFileBtn = Button(frame2, text="Choose Message File", command= lambda:openFile("MS"))

frame3 = LabelFrame(root, text="Message Received")
labelMess2 = Label(frame3, text="Message Received")
entryMess2 = Text(frame3,height=10,width=30,borderwidth=3)
# entryMess2 = Entry(frame3, width=30,borderwidth=5)
labelSignature2 = Label(frame3, text="Digital Signature Received")
entrySignature2 = Text(frame3,height=10,width=30, borderwidth=3)
# entrySignature2 = Entry(frame3, width=30,borderwidth=5)
verifieSignatureBtn = Button(frame3, text="Signature Verification", command= signatureVerification)
chooseFileBtn2 = Button(frame3, text="Choose Message File", command= lambda: openFile("MR"))
chooseFileBtn3 = Button(frame3, text="Choose Signature File", command= lambda: openFile("DS"))

frame4 = LabelFrame(root, text="Puclic Key")
labelPubKeyP = Label(frame4, text="Public key (p)")
entryPubKeyP = Entry(frame4, width=30,borderwidth=3)
labelPubKeyAlpha = Label(frame4, text="Public key (Alpha)")
entryPubKeyAlpha = Entry(frame4, width=30,borderwidth=3)
labelPubKeyBeta = Label(frame4, text="Public key (Beta)")
entryPubKeyBeta = Entry(frame4, width=30,borderwidth=3)
enterPubKey = Button(frame4, text="Enter Puclic Key", command=enterPublicKey)

# Hiển thị lên màn hình

frame1.grid(row=0, column=0, padx=5, pady=5)
labelParameter.grid(row=0,column=0,padx=5,pady=5)
entryParameter.grid(row=1,column=0,padx=5,pady=5)
labelPubKey.grid(row=2,column=0,padx=5,pady=5)
entryPubKey.grid(row=3,column=0,padx=5,pady=5)
labelPriKey.grid(row=4,column=0,padx=5,pady=5)
entryPriKey.grid(row=5,column=0,padx=5,pady=5)
labelPriK.grid(row=6,column=0,padx=5,pady=5)
entryPriK.grid(row=7,column=0,padx=5,pady=5)
generateKeyBtn.grid(row=8,column=0,padx=5,pady=5)

frame2.grid(row=0, column=1, padx=5, pady=5)
labelMess.grid(row=0,column=1,padx=5,pady=5)
entryMess.grid(row=1,column=1,padx=5,pady=5)
labelSignature.grid(row=2,column=1,padx=5,pady=5)
entrySignature.grid(row=3,column=1,padx=5,pady=5)
generateSignatureBtn.grid(row=4,column=1,padx=5,pady=5)
# enterBtn.grid(row=1,column=2)
chooseFileBtn.grid(row=1,column=2,padx=5)

frame3.grid(row=0, column=3, padx=10, pady=5)
labelMess2.grid(row=0,column=3,padx=5,pady=5)
entryMess2.grid(row=1,column=3,padx=5,pady=5)
labelSignature2.grid(row=2,column=3,padx=5,pady=5)
entrySignature2.grid(row=3,column=3,padx=5,pady=5)
verifieSignatureBtn.grid(row=4,column=3,padx=5,pady=5)
chooseFileBtn2.grid(row=1,column=4, padx=5)
chooseFileBtn3.grid(row=3,column=4)

frame4.grid(row=0, column=4)
labelPubKeyP.grid(row=0,column=4,padx=5,pady=5)
entryPubKeyP.grid(row=1,column=4,padx=5,pady=5)
labelPubKeyAlpha.grid(row=2,column=4,padx=5,pady=5)
entryPubKeyAlpha.grid(row=3,column=4,padx=5,pady=5)
labelPubKeyBeta.grid(row=4,column=4,padx=5,pady=5)
entryPubKeyBeta.grid(row=5,column=4,padx=5,pady=5)
enterPubKey.grid(row=6,column=4,padx=5,pady=5)

root.mainloop()
